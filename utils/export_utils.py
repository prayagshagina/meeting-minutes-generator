import os
import logging

logger = logging.getLogger(__name__)

# Try importing optional libraries
_has_reportlab = False
_has_docx = False
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    _has_reportlab = True
except Exception:
    logger.info("reportlab not installed — PDF export disabled.")

try:
    from docx import Document
    _has_docx = True
except Exception:
    logger.info("python-docx not installed — DOCX export disabled.")


class Exporter:
    def __init__(self, export_dir: str):
        self.export_dir = export_dir
        os.makedirs(self.export_dir, exist_ok=True)

    def export_txt(self, basename: str, title: str, date: str, structured: dict, transcript: str) -> str:
        filename = os.path.join(self.export_dir, f"{basename}.txt")
        with open(filename, "w", encoding="utf-8") as f:
            f.write(f"Title: {title}\nDate: {date}\n\n")
            f.write("SUMMARY:\n")
            f.write(structured.get("summary", "") + "\n\n")
            f.write("KEY POINTS:\n")
            for kp in structured.get("key_points", []):
                f.write("- " + kp + "\n")
            f.write("\nACTION ITEMS:\n")
            for ai in structured.get("action_items", []):
                f.write("- " + ai + "\n")
            f.write("\nDEADLINES:\n")
            for d in structured.get("deadlines", []):
                f.write("- " + d + "\n")
            f.write("\nTRANSCRIPT:\n")
            f.write(transcript + "\n")
        logger.info(f"Wrote TXT export: {filename}")
        return filename

    def export_pdf(self, basename: str, title: str, date: str, structured: dict, transcript: str):
        if not _has_reportlab:
            logger.info("Skipping PDF export (reportlab missing).")
            return None
        filename = os.path.join(self.export_dir, f"{basename}.pdf")
        try:
            c = canvas.Canvas(filename, pagesize=A4)
            width, height = A4
            y = height - 50
            c.setFont("Helvetica-Bold", 14)
            c.drawString(40, y, f"Meeting Minutes - {title}")
            y -= 20
            c.setFont("Helvetica", 10)
            c.drawString(40, y, f"Date: {date}")
            y -= 24

            def draw_block(title_text, content_lines):
                nonlocal y
                if y < 120:
                    c.showPage()
                    y = height - 50
                c.setFont("Helvetica-Bold", 12)
                c.drawString(40, y, title_text)
                y -= 16
                c.setFont("Helvetica", 10)
                for line in content_lines:
                    # wrap naive: break every ~90 chars
                    while len(line) > 90:
                        c.drawString(48, y, line[:90])
                        line = line[90:]
                        y -= 14
                        if y < 80:
                            c.showPage()
                            y = height - 50
                    c.drawString(48, y, line)
                    y -= 14
                    if y < 80:
                        c.showPage()
                        y = height - 50
                y -= 8

            draw_block("Summary", [structured.get("summary", "")])
            draw_block("Key Points", structured.get("key_points", []))
            draw_block("Action Items", structured.get("action_items", []))
            draw_block("Deadlines", structured.get("deadlines", []))
            draw_block("Transcript (begin)", [transcript[:4000] + ("..." if len(transcript) > 4000 else "")])

            c.save()
            logger.info(f"Wrote PDF export: {filename}")
            return filename
        except Exception:
            logger.exception("PDF export failed.")
            return None

    def export_docx(self, basename: str, title: str, date: str, structured: dict, transcript: str):
        if not _has_docx:
            logger.info("Skipping DOCX export (python-docx missing).")
            return None
        filename = os.path.join(self.export_dir, f"{basename}.docx")
        try:
            doc = Document()
            doc.add_heading(f"Meeting Minutes - {title}", level=1)
            doc.add_paragraph(f"Date: {date}")
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(structured.get("summary", ""))
            if structured.get("key_points"):
                doc.add_heading("Key Points", level=2)
                for kp in structured.get("key_points", []):
                    doc.add_paragraph(kp, style='List Bullet')
            if structured.get("action_items"):
                doc.add_heading("Action Items", level=2)
                for ai in structured.get("action_items", []):
                    doc.add_paragraph(ai, style='List Bullet')
            if structured.get("deadlines"):
                doc.add_heading("Deadlines", level=2)
                for d in structured.get("deadlines", []):
                    doc.add_paragraph(d, style='List Bullet')

            doc.add_heading("Transcript", level=2)
            doc.add_paragraph(transcript)
            doc.save(filename)
            logger.info(f"Wrote DOCX export: {filename}")
            return filename
        except Exception:
            logger.exception("DOCX export failed.")
            return None
